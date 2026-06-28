"""Generate a Korean lab introduction docx (연구실 소개서)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_korean_font(run, font_name='맑은 고딕', size=10, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_heading(doc, text, level=1):
    h = doc.add_paragraph()
    run = h.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_korean_font(run, size=sizes.get(level, 12), bold=True)
    if level == 1:
        run.font.color.rgb = RGBColor(0x00, 0x6E, 0x5A)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x25)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)


def add_kv(doc, key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rk = p.add_run(f'  • {key}: ')
    set_korean_font(rk, size=10, bold=True)
    rv = p.add_run(value)
    set_korean_font(rv, size=10)


def add_body(doc, text, indent_cm=0, bullet=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    prefix = '  • ' if bullet else ''
    run = p.add_run(prefix + text)
    set_korean_font(run, size=10)


def add_pub(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    run = p.add_run(f'[{n}] {text}')
    set_korean_font(run, size=9.5)


doc = Document()

# Set default style
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('IO-LAB 연구실 소개')
set_korean_font(tr, size=22, bold=True)
tr.font.color.rgb = RGBColor(0x00, 0x6E, 0x5A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run('지능형 최적화 연구실 · Intelligent Optimization Laboratory')
set_korean_font(sr, size=12)
sr.font.color.rgb = RGBColor(0x4A, 0x4A, 0x5A)

doc.add_paragraph()  # spacer

# 1. 연구실 정보
add_heading(doc, '1. 연구실 기본 정보', 1)
add_kv(doc, '연구실 이름', '지능형 최적화 연구실 (IO-LAB, Intelligent Optimization Laboratory)')
add_kv(doc, '소속', '경희대학교 공과대학 산업경영공학과')
add_kv(doc, '연구책임자', 'Ho Young Jeong (정호영), Ph.D., 조교수')
add_kv(doc, '연구실 위치', '경희대학교 국제캠퍼스 공학관 504호')
add_kv(doc, '주소', '경기도 용인시 기흥구 덕영대로 1732')
add_kv(doc, '교수실', '경희대학교 국제캠퍼스 우정원 6034호')
add_kv(doc, '전화', '031-201-2430 (교수실)')
add_kv(doc, '이메일', 'ghy27@khu.ac.kr')
add_kv(doc, '홈페이지', 'https://khu-iolab.github.io')
add_kv(doc, 'YouTube', 'https://www.youtube.com/@IO-LAB-khu')
add_kv(doc, 'Instagram', 'https://www.instagram.com/khu_iolab/')

# 2. 연구실 소개
add_heading(doc, '2. 연구실 소개', 1)
add_body(doc,
    '지능형 최적화 연구실(IO-LAB)은 경희대학교 산업경영공학과에 소속된 연구실로, '
    '운영과학(Operations Research)과 최적화 기법을 핵심 도구로 활용하여 물류·제조·생산·통합 시스템의 미래를 개척합니다.')
add_body(doc,
    '본 연구실은 도시 물류, 드론 및 자율 모빌리티 배송, 공급망 회복력, 반도체 팹 자동화 물류, '
    'V2G 기반 에너지-물류 통합 운영, 메타러닝·강화학습·양자 컴퓨팅 기반 차세대 최적화 알고리즘 개발 등 '
    '폭넓은 주제를 다룹니다.')
add_body(doc,
    '수리적 최적화 모델링, 휴리스틱·메타휴리스틱, 시뮬레이션 기반 최적화, 머신러닝·딥러닝과의 융합을 '
    '통해 이론적 기여와 산업적 적용을 함께 추구합니다.')

# 3. 연구분야
add_heading(doc, '3. 연구분야', 1)

add_heading(doc, '3.1 물류 최적화 (Logistics Optimization)', 2)
add_body(doc,
    '도시 물류, 드론 배송, 지하 물류 네트워크, 비행 창고(Flying Warehouse) 등 복잡한 물류 시스템의 '
    '최적 운영 전략을 연구합니다. 차량-드론 복합 배송, 다중 비행 고도(multi-flight level) 드론 경로 설계, '
    '지하 운송망과 연계된 도시 물류 최적화 등을 다룹니다.')
add_body(doc,
    '주요 키워드: 차량 경로 문제(VRP), 드론 배송, 라스트마일, 비행 창고, 지하 물류, 공급망 설계',
    bullet=False)

add_heading(doc, '3.2 네트워크 최적화 (Network Optimization)', 2)
add_body(doc,
    '기후변화 불확실성을 고려한 공급망(supply chain network) 설계 및 적응형 운영 전략, '
    '시뮬레이션 기반 장애 전파·복구 전략과 네트워크 회복력(resilience) 분석, '
    'V2G 기반 EV 물류·배전망 통합 운영 등 복잡한 네트워크 시스템의 설계·운영·강건성을 연구합니다.')
add_body(doc,
    '주요 키워드: 공급망 네트워크, 기후변화 대응, 군사 네트워크, 네트워크 회복력, '
    '시뮬레이션 최적화, V2G/G2V, EV 물류, 배전망 최적화')

add_heading(doc, '3.3 생산 최적화 (Production Optimization)', 2)
add_body(doc,
    '반도체 팹 등 첨단 제조 시설의 자동화 물류 시스템(OHT, AGV, AMR) 배차·경로 최적화·교통 제어 전략을 '
    '시뮬레이션 기반으로 비교·최적화하고, 열간 단조·스탬핑 등 제조 공정의 파라미터 최적화와 '
    '머신러닝 통합 FEM 시뮬레이션을 결합하여 생산 처리량과 효율을 극대화하는 연구를 수행합니다.')
add_body(doc,
    '주요 키워드: 반도체 팹 자동화(OHT/AGV/AMR), 열간 단조, 스탬핑, 유한요소법(FEM), '
    '공정 최적화, 품질 제어, 지능형 제조')

add_heading(doc, '3.4 차세대 최적화 (Next-Gen Optimization)', 2)
add_body(doc,
    '메타러닝(meta-learning), 강화학습, 신경망 기반 해법을 활용한 조합 최적화 문제 해결 및 '
    '양자 컴퓨팅(quantum computing) 기반 알고리즘 개발 등 미래 지향적 최적화 기법을 연구합니다. '
    '인공지능과 운영과학의 융합을 통해 복잡계 문제에 도전합니다.')
add_body(doc,
    '주요 키워드: 메타러닝, 강화학습, 딥러닝, 양자 컴퓨팅, 조합 최적화, 인공지능')

# 4. 논문
add_heading(doc, '4. 주요 논문 (Selected Publications)', 1)

add_heading(doc, '4.1 국제 학술지 (SCIE)', 2)
pubs_intl = [
    'Jeong, H.Y., Song, B.D. (2025). "Optimization of Urban Logistics with Multi-Modal Systems: A Comprehensive Study of the Airship-Vehicle Routing Problem." Transportation Research Part E: Logistics and Transportation Review, 204, p.104415. [JCR 1.4%]',
    'Jeong, H.Y., Song, B.D. (2025). "Meta-Learning-Based Adaptive Operator Selection for Traveling Salesman Problem." Applied Soft Computing, p.113930. [JCR 9.1%]',
    'Jeong, H.Y., Song, B.D. (2025). "Optimizing Urban Logistics: Vehicle Routing Problem with Underground Transportation." IEEE Transactions on Intelligent Transportation Systems. [JCR 8%]',
    'Kim, Y., Jeong, H.Y.*, Lee, S. (2024). "Drone delivery problem with multi-flight level: Machine learning based solution approach." Computers & Industrial Engineering, 197, p.110565. [JCR 15%]',
    'Guo, F., Jeong, H.Y., Park, D., Kim, G., Sung, B., Kim, N. (2024). "Numerical optimization of variable blank holder force trajectories in stamping process for multi-defect reduction." Materials, 17(11), p.2578.',
    'Guo, F., Jeong, H.Y., Park, D., Sung, B., Kim, N.* (2024). "Numerical multi-objective optimization of segmented and variable blank holder force trajectories in deep drawing based on DNN-GA-MCS strategy." The International Journal of Advanced Manufacturing Technology.',
    'Kim, Y., Jeong, H.Y., Park, J., Kim, K., Kwon, H., et al. (2023). "Optimizing process parameters for hot forging of Ti-6242 alloy: ML and FEM simulation." Journal of Materials Research and Technology, 27, pp.8228–8243.',
    'Kim, B., Jeong, H.Y.*, Lee, S. (2023). "Two-echelon collaborative routing problem with heterogeneous crowd-shippers." Computers & Operations Research, 160, p.106389.',
    'Jeong, H.Y., Choi, C.* (2023). "Adaptive supply chain system design for fruit crops under climate change." Systems, 11(10), p.514.',
    'Jeong, H.Y., Lee, S.* (2023). "Drone routing problem with truck: Optimization and quantitative analysis." Expert Systems with Applications, 227, p.120260.',
    'Park, J., Kim, Y., Jeong, H.Y., et al. (2023). "Cogging process design of M50 bearing steel for billet quality." Journal of Materials Research and Technology, 26, pp.5576–5593.',
    'Jeong, H.Y.*, Kim, Y., Lee, S., Moreland, J., Zhou, C. (2023). "Disruption propagation and repair response in interdependent systems." Simulation Modelling Practice and Theory, 124, p.102730.',
    'Jeong, H.Y., Park, J., Kim, Y., Shin, S.Y., Kim, N.* (2023). "Processing parameters optimization in hot forging of AISI 4340 steel using instability map and RL." Journal of Materials Research and Technology, 23, pp.1995–2009.',
    'Jeong, H.Y., Song, B.D.*, Lee, S. (2022). "Optimal scheduling for multi-flying warehouse: Amazon airborne fulfillment center." Transportation Research Part C: Emerging Technologies, 143, p.103831.',
    'Jeong, H.Y., Song, B.D.*, Lee, S. (2020). "The flying warehouse delivery system." IEEE Transactions on Intelligent Transportation Systems, 22(12).',
    'Jeong, H.Y., David, J.Y., Min, B.C., Lee, S.* (2020). "The humanitarian flying warehouse." Transportation Research Part E: Logistics and Transportation Review, 136, p.101901. [JCR 1.4%]',
    'Jeong, H.Y., Song, B.D.*, Lee, S. (2019). "Truck-drone hybrid delivery routing: Payload-energy dependency and No-Fly zones." International Journal of Production Economics, 214, pp.220–233. [JCR 5%]',
]
for i, t in enumerate(pubs_intl, 1):
    add_pub(doc, i, t)

add_heading(doc, '4.2 학회 발표 및 학술지', 2)
pubs_conf = [
    'Kim, B., Jeong, H.Y., Lee, S. "Last-Mile Delivery Capacity Building in Pandemics Utilizing Community Resources." SSRN 4168513.',
    'Kim, Y., Jung, H., Lee, S. (2021). IFIP International Conference on Advances in Production Management Systems, Cham: Springer, pp. 43–51.',
    'Jeong, H.Y., Lee, S. (2021). Advances in Production Management Systems. IFIP WG 5.7 International Conference, Springer, pp. 33–42.',
    'Song, B.D., Jun, S., Jung, H.Y., Lee, S. (2019). Procedia Manufacturing, 39, pp.300–306.',
    'Jeong, H.Y., Lee, S. (2019). Procedia Manufacturing, 39, pp.307–313.',
    'Jeong, H.Y., Lee, S. "Airship-based drone delivery system: Quantitative approach for managerial and operational decisions." INFORMS Annual Meeting, Seattle, US, 2019.',
    'Jeong, H.Y., Lee, S. ICPR25, 2019. "Optimization of Vehicle-Carrier Routing Mathematical Model and Comparison with Related Variants."',
    'Jeong, H.Y., Lee, S. INFORMS Annual Meeting, Phoenix, US, 2018. "Vehicle-Carrier Routing Problem."',
    'Jeong, H.Y., Lee, S. IISE Annual Meeting, Orlando, US, 2018. "Scheduling Hybrid Delivery System of Truck and Drone."',
]
for i, t in enumerate(pubs_conf, 1):
    add_pub(doc, i, t)

add_heading(doc, '4.3 Working Papers', 2)
pubs_wp = [
    '(Major Revision) Choi, C.H., Jeong, H.Y.*. "Integrated long-term capacity expansion and operational scheduling for regional power grids: A MISO case study."',
    '(Under Review) Jeong, H.Y. "Simulation-Based Hierarchical Optimization for Urban Drone Delivery in Multi-Altitude 3D Airspace."',
    '(Under Review) Jeong, H.Y. "Dispatch Strategy Comparison for Overhead Hoist Transport Systems: From State-Aware Dispatching to Predictive Repositioning Across Layout Scales."',
]
for i, t in enumerate(pubs_wp, 1):
    add_pub(doc, i, t)

add_heading(doc, '4.4 특허', 2)
add_pub(doc, 1, 'Kim, N., Jeong, H.Y. Korean Patent No. 10-2023-0071207.')

# Footer note
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
nr = note.add_run('※ 최신 정보는 연구실 홈페이지(https://khu-iolab.github.io)에서 확인하실 수 있습니다.')
set_korean_font(nr, size=8.5)
nr.font.color.rgb = RGBColor(0x6A, 0x6A, 0x7A)
nr.italic = True

# Save
out = 'IO-LAB_연구실_소개서.docx'
doc.save(out)
print(f'Saved: {out}')
